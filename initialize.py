"""
このファイルは、最初の画面読み込み時にのみ実行される初期化処理が記述されたファイルです。
"""

############################################################
# ライブラリの読み込み
############################################################
import os
import logging
from logging.handlers import TimedRotatingFileHandler
from uuid import uuid4
import sys
import unicodedata
import csv
from dotenv import load_dotenv
import streamlit as st
from docx import Document
from langchain_community.document_loaders import WebBaseLoader
from langchain.text_splitter import CharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.schema import Document as LangChainDocument
import constants as ct


############################################################
# 設定関連
############################################################
# 「.env」ファイルで定義した環境変数の読み込み
load_dotenv()


############################################################
# 関数定義
############################################################

def initialize():
    """
    画面読み込み時に実行する初期化処理
    """
    try:
        # 初期化データの用意
        initialize_session_state()
        print("[DEBUG] Session state initialized")
        
        # ログ出力用にセッションIDを生成
        initialize_session_id()
        print("[DEBUG] Session ID initialized")
        
        # ログ出力の設定
        initialize_logger()
        print("[DEBUG] Logger initialized")
        
        # RAGのRetrieverを作成
        initialize_retriever()
        print("[DEBUG] Retriever initialized")
        
    except Exception as e:
        print(f"[ERROR] Initialization failed: {str(e)}")
        raise


def initialize_logger():
    """
    ログ出力の設定
    """
    # 指定のログフォルダが存在すれば読み込み、存在しなければ新規作成
    os.makedirs(ct.LOG_DIR_PATH, exist_ok=True)
    
    # 引数に指定した名前のロガー（ログを記録するオブジェクト）を取得
    # 再度別の箇所で呼び出した場合、すでに同じ名前のロガーが存在していれば読み込む
    logger = logging.getLogger(ct.LOGGER_NAME)

    # すでにロガーにハンドラー（ログの出力先を制御するもの）が設定されている場合、同じログ出力が複数回行われないよう処理を中断する
    if logger.hasHandlers():
        return

    # 1日単位でログファイルの中身をリセットし、切り替える設定
    log_handler = TimedRotatingFileHandler(
        os.path.join(ct.LOG_DIR_PATH, ct.LOG_FILE),
        when="D",
        encoding="utf8"
    )
    # 出力するログメッセージのフォーマット定義
    # - 「levelname」: ログの重要度（INFO, WARNING, ERRORなど）
    # - 「asctime」: ログのタイムスタンプ（いつ記録されたか）
    # - 「lineno」: ログが出力されたファイルの行番号
    # - 「funcName」: ログが出力された関数名
    # - 「session_id」: セッションID（誰のアプリ操作か分かるように）
    # - 「message」: ログメッセージ
    formatter = logging.Formatter(
        f"[%(levelname)s] %(asctime)s line %(lineno)s, in %(funcName)s, session_id={st.session_state.session_id}: %(message)s"
    )

    # 定義したフォーマッターの適用
    log_handler.setFormatter(formatter)

    # ログレベルを「INFO」に設定
    logger.setLevel(logging.INFO)

    # 作成したハンドラー（ログ出力先を制御するオブジェクト）を、
    # ロガー（ログメッセージを実際に生成するオブジェクト）に追加してログ出力の最終設定
    logger.addHandler(log_handler)


def initialize_session_id():
    """
    セッションIDの作成
    """
    if "session_id" not in st.session_state:
        # ランダムな文字列（セッションID）を、ログ出力用に作成
        st.session_state.session_id = uuid4().hex


def initialize_retriever():
    """
    画面読み込み時にRAGのRetriever（ベクターストアから検索するオブジェクト）を作成
    """
    # ロガーを読み込むことで、後続の処理中に発生したエラーなどがログファイルに記録される
    logger = logging.getLogger(ct.LOGGER_NAME)

    # すでにRetrieverが作成済みの場合、後続の処理を中断
    # デバッグ用：強制的にリセットするためのコメントアウト
    # if "retriever" in st.session_state:
    #     return
    
    # 改善のため、毎回Retrieverを再作成
    print("[DEBUG] Force recreating retriever for CSV improvements")
    
    # 既存のRetrieverを削除
    if "retriever" in st.session_state:
        del st.session_state.retriever
        print("[DEBUG] Cleared existing retriever")
    
    # OpenAI API キーの確認
    openai_api_key = os.getenv("OPENAI_API_KEY")
    if not openai_api_key:
        raise ValueError("OPENAI_API_KEY environment variable is not set")
    
    try:
        # RAGの参照先となるデータソースの読み込み
        docs_all = load_data_sources()
        
        if not docs_all:
            raise ValueError("No documents found to create embeddings")

        # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
        for doc in docs_all:
            doc.page_content = adjust_string(doc.page_content)
            for key in doc.metadata:
                doc.metadata[key] = adjust_string(doc.metadata[key])
        
        # 埋め込みモデルの用意
        embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
        
        # チャンク分割用のオブジェクトを作成
        text_splitter = CharacterTextSplitter(
            chunk_size=ct.CHUNK_SIZE,
            chunk_overlap=ct.CHUNK_OVERLAP,
            separator="\n"
        )

        # チャンク分割を実施
        splitted_docs = text_splitter.split_documents(docs_all)
        
        if not splitted_docs:
            raise ValueError("No documents after splitting")

        # 各チャンクの情報をデバッグ出力
        hr_chunks = [doc for doc in splitted_docs if "人事部" in doc.page_content]
        print(f"[DEBUG] HR-related chunks: {len(hr_chunks)}")
        
        # ベクターストアの作成
        db = FAISS.from_documents(splitted_docs, embeddings)

        # ベクターストアを検索するRetrieverの作成
        st.session_state.retriever = db.as_retriever(search_kwargs={"k": ct.RETRIEVER_SEARCH_COUNT})
        
        logger.info(f"Retriever created successfully with {len(splitted_docs)} document chunks")
        print(f"[DEBUG] Retriever created successfully with {len(splitted_docs)} document chunks")
        
        # デバッグ用：人事部関連の検索テスト
        test_docs = st.session_state.retriever.get_relevant_documents("人事部に所属している従業員")
        print(f"[DEBUG] Test search returned {len(test_docs)} documents")
        for i, doc in enumerate(test_docs[:3]):
            print(f"[DEBUG] Test doc {i+1}: {doc.page_content[:200]}...")
        
    except Exception as e:
        logger.error(f"Error in initialize_retriever: {str(e)}")
        print(f"[ERROR] Error in initialize_retriever: {str(e)}")
        raise


def initialize_session_state():
    """
    初期化データの用意
    """
    if "messages" not in st.session_state:
        # 「表示用」の会話ログを順次格納するリストを用意
        st.session_state.messages = []
        # 「LLMとのやりとり用」の会話ログを順次格納するリストを用意
        st.session_state.chat_history = []


def load_data_sources():
    """
    RAGの参照先となるデータソースの読み込み

    Returns:
        読み込んだ通常データソース
    """
    logger = logging.getLogger(ct.LOGGER_NAME)
    
    # データソースを格納する用のリスト
    docs_all = []
    
    try:
        # データフォルダの存在確認
        if not os.path.exists(ct.RAG_TOP_FOLDER_PATH):
            logger.warning(f"Data folder not found: {ct.RAG_TOP_FOLDER_PATH}")
            print(f"[WARNING] Data folder not found: {ct.RAG_TOP_FOLDER_PATH}")
        else:
            logger.info(f"Data folder found: {ct.RAG_TOP_FOLDER_PATH}")
            print(f"[INFO] Data folder found: {ct.RAG_TOP_FOLDER_PATH}")
        
        # ファイル読み込みの実行（渡した各リストにデータが格納される）
        recursive_file_check(ct.RAG_TOP_FOLDER_PATH, docs_all)
        logger.info(f"Loaded {len(docs_all)} documents from local files")
        print(f"[INFO] Loaded {len(docs_all)} documents from local files")

        web_docs_all = []
        # ファイルとは別に、指定のWebページ内のデータも読み込み
        # 読み込み対象のWebページ一覧に対して処理
        for web_url in ct.WEB_URL_LOAD_TARGETS:
            try:
                # 指定のWebページを読み込み
                loader = WebBaseLoader(web_url)
                web_docs = loader.load()
                # for文の外のリストに読み込んだデータソースを追加
                web_docs_all.extend(web_docs)
                logger.info(f"Loaded {len(web_docs)} documents from {web_url}")
                print(f"[INFO] Loaded {len(web_docs)} documents from {web_url}")
            except Exception as e:
                logger.warning(f"Failed to load web content from {web_url}: {str(e)}")
                print(f"[WARNING] Failed to load web content from {web_url}: {str(e)}")
                continue
        
        # 通常読み込みのデータソースにWebページのデータを追加
        docs_all.extend(web_docs_all)
        logger.info(f"Total documents loaded: {len(docs_all)}")
        print(f"[INFO] Total documents loaded: {len(docs_all)}")
        
    except Exception as e:
        logger.error(f"Error loading data sources: {str(e)}")
        print(f"[ERROR] Error loading data sources: {str(e)}")
        raise

    return docs_all


def recursive_file_check(path, docs_all):
    """
    RAGの参照先となるデータソースの読み込み

    Args:
        path: 読み込み対象のファイル/フォルダのパス
        docs_all: データソースを格納する用のリスト
    """
    try:
        # パスがフォルダかどうかを確認
        if os.path.isdir(path):
            # フォルダの場合、フォルダ内のファイル/フォルダ名の一覧を取得
            files = os.listdir(path)
            print(f"[DEBUG] Checking directory: {path}, found {len(files)} items")
            # 各ファイル/フォルダに対して処理
            for file in files:
                # ファイル/フォルダ名だけでなく、フルパスを取得
                full_path = os.path.join(path, file)
                # フルパスを渡し、再帰的にファイル読み込みの関数を実行
                recursive_file_check(full_path, docs_all)
        else:
            # パスがファイルの場合、ファイル読み込み
            file_load(path, docs_all)
    except Exception as e:
        print(f"[ERROR] Error in recursive_file_check for {path}: {str(e)}")
        raise


def file_load(path, docs_all):
    """
    ファイル内のデータ読み込み

    Args:
        path: ファイルパス
        docs_all: データソースを格納する用のリスト
    """
    try:
        # ファイルの拡張子を取得
        file_extension = os.path.splitext(path)[1]
        # ファイル名（拡張子を含む）を取得
        file_name = os.path.basename(path)

        # 想定していたファイル形式の場合のみ読み込む
        if file_extension in ct.SUPPORTED_EXTENSIONS:
            print(f"[DEBUG] Loading file: {path}")
            
            # CSVファイルの場合は特別な処理を行う
            if file_extension == ".csv":
                csv_docs = load_csv_as_unified_document(path)
                docs_all.extend(csv_docs)
                print(f"[DEBUG] Loaded {len(csv_docs)} unified documents from CSV: {path}")
            else:
                # 他のファイル形式は通常通り処理
                loader = ct.SUPPORTED_EXTENSIONS[file_extension](path)
                docs = loader.load()
                docs_all.extend(docs)
                print(f"[DEBUG] Loaded {len(docs)} documents from {path}")
        else:
            print(f"[DEBUG] Skipping unsupported file: {path} (extension: {file_extension})")
    except Exception as e:
        print(f"[ERROR] Error loading file {path}: {str(e)}")
        # ファイル読み込みエラーは警告のみで処理を継続
        pass


def load_csv_as_unified_document(csv_path):
    """
    CSVファイルを統合されたドキュメントとして読み込む
    
    Args:
        csv_path: CSVファイルのパス
    
    Returns:
        統合されたドキュメントのリスト
    """
    try:
        print(f"[DEBUG] Processing CSV file: {csv_path}")
        
        # CSVファイルを読み込む
        with open(csv_path, 'r', encoding='utf-8') as file:
            reader = csv.DictReader(file)
            rows = list(reader)
        
        print(f"[DEBUG] CSV rows loaded: {len(rows)}")
        
        if not rows:
            return []
        
        # 統合されたドキュメントを作成
        unified_docs = []
        
        # 人事部専用の大きな統合ドキュメントを作成（最優先）
        hr_employees = [row for row in rows if row.get('部署') == '人事部']
        print(f"[DEBUG] HR employees found: {len(hr_employees)}")
        
        if hr_employees:
            # 人事部の大きな統合ドキュメント
            hr_text = "■ 人事部従業員完全一覧\n"
            hr_text += f"人事部所属従業員総数: {len(hr_employees)}名\n\n"
            hr_text += "人事部に所属している従業員情報一覧:\n\n"
            
            for i, emp in enumerate(hr_employees, 1):
                name = emp.get('氏名（フルネーム）', '不明')
                emp_id = emp.get('社員ID', '不明')
                position = emp.get('役職', '不明')
                emp_type = emp.get('従業員区分', '不明')
                age = emp.get('年齢', '不明')
                join_date = emp.get('入社日', '不明')
                skills = emp.get('スキルセット', '不明')
                qualifications = emp.get('保有資格', '不明')
                university = emp.get('大学名', '不明')
                faculty = emp.get('学部・学科', '不明')
                email = emp.get('メールアドレス', '不明')
                
                hr_text += f"【{i}】人事部従業員\n"
                hr_text += f"氏名: {name}\n"
                hr_text += f"社員ID: {emp_id}\n"
                hr_text += f"部署: 人事部\n"
                hr_text += f"役職: {position}\n"
                hr_text += f"従業員区分: {emp_type}\n"
                hr_text += f"年齢: {age}歳\n"
                hr_text += f"入社日: {join_date}\n"
                hr_text += f"スキルセット: {skills}\n"
                hr_text += f"保有資格: {qualifications}\n"
                hr_text += f"大学: {university} {faculty}\n"
                hr_text += f"メールアドレス: {email}\n"
                hr_text += f"人事部所属: はい\n"
                hr_text += f"人事部門: 人事部\n\n"
            
            # 人事部の名前リスト
            hr_text += "人事部従業員名簿:\n"
            for i, emp in enumerate(hr_employees, 1):
                hr_text += f"{i}. {emp.get('氏名（フルネーム）', '不明')} ({emp.get('社員ID', '不明')})\n"
            
            hr_doc = LangChainDocument(
                page_content=hr_text,
                metadata={
                    "source": csv_path,
                    "department": "人事部",
                    "employee_count": len(hr_employees),
                    "type": "hr_complete_list",
                    "file_name": os.path.basename(csv_path),
                    "description": "人事部全従業員の完全な一覧"
                }
            )
            unified_docs.append(hr_doc)
            print(f"[DEBUG] Created HR complete document with {len(hr_employees)} employees")
        
        # 部署別の統合ドキュメントを作成
        department_groups = {}
        for row in rows:
            department = row.get('部署', '不明')
            if department not in department_groups:
                department_groups[department] = []
            department_groups[department].append(row)
        
        print(f"[DEBUG] Department groups: {list(department_groups.keys())}")
        
        for department, employees in department_groups.items():
            # 部署別統合ドキュメント
            dept_text = f"■ {department}の従業員一覧\n"
            dept_text += f"部署名: {department}\n"
            dept_text += f"所属人数: {len(employees)}名\n\n"
            
            for i, emp in enumerate(employees, 1):
                name = emp.get('氏名（フルネーム）', '不明')
                emp_id = emp.get('社員ID', '不明')
                position = emp.get('役職', '不明')
                skills = emp.get('スキルセット', '不明')
                
                dept_text += f"{i}. {name} ({emp_id})\n"
                dept_text += f"   部署: {department}\n"
                dept_text += f"   役職: {position}\n"
                dept_text += f"   スキルセット: {skills}\n\n"
            
            dept_doc = LangChainDocument(
                page_content=dept_text,
                metadata={
                    "source": csv_path,
                    "department": department,
                    "employee_count": len(employees),
                    "type": "department_summary",
                    "file_name": os.path.basename(csv_path)
                }
            )
            unified_docs.append(dept_doc)
        
        # 全社員の統合ドキュメント
        all_text = "■ 全社員名簿\n"
        all_text += f"全社員数: {len(rows)}名\n\n"
        
        # 部署別サマリー
        all_text += "部署別従業員数:\n"
        for dept, employees in department_groups.items():
            all_text += f"- {dept}: {len(employees)}名\n"
        all_text += "\n"
        
        # 全従業員リスト
        for i, row in enumerate(rows, 1):
            all_text += f"{i}. {row.get('氏名（フルネーム）', '不明')} ({row.get('社員ID', '不明')})\n"
            all_text += f"   部署: {row.get('部署', '不明')}\n"
            all_text += f"   役職: {row.get('役職', '不明')}\n\n"
        
        all_doc = LangChainDocument(
            page_content=all_text,
            metadata={
                "source": csv_path,
                "department": "全社",
                "employee_count": len(rows),
                "type": "all_employees_summary",
                "file_name": os.path.basename(csv_path)
            }
        )
        unified_docs.append(all_doc)
        
        print(f"[DEBUG] Total unified documents created: {len(unified_docs)}")
        
        return unified_docs
        
    except Exception as e:
        print(f"[ERROR] Error in load_csv_as_unified_document: {str(e)}")
        # エラーが発生した場合は通常のCSVLoader処理にフォールバック
        try:
            print(f"[DEBUG] Falling back to normal CSV loader")
            loader = ct.SUPPORTED_EXTENSIONS[".csv"](csv_path)
            return loader.load()
        except:
            return []


def adjust_string(s):
    """
    Windows環境でRAGが正常動作するよう調整
    
    Args:
        s: 調整を行う文字列
    
    Returns:
        調整を行った文字列
    """
    # 調整対象は文字列のみ
    if type(s) is not str:
        return s

    # OSがWindowsの場合、Unicode正規化と、cp932（Windows用の文字コード）で表現できない文字を除去
    if sys.platform.startswith("win"):
        s = unicodedata.normalize('NFC', s)
        s = s.encode("cp932", "ignore").decode("cp932")
        return s
    
    # OSがWindows以外の場合はそのまま返す
    return s